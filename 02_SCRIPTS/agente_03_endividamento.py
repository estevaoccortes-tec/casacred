#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
AGENTE_03 — Endividamento (multi-formato)

Lê 1 arquivo de endividamento (XLSX/CSV/PDF/DOCX/PNG/JPG) e gera um XLSX
com 1 aba (Planilha1) e 5 colunas FIXAS:
1) Instituição Financeira
2) Tipo
3) Linha de Crédito
4) Limite Utilizado
5) CNPJ

Regra do valor (adaptada pro seu cenário real):
- Aceitar valor vindo de cabeçalhos que contenham:
  * "utiliz" (saldo utilizado, limite utilizado, valor utilizado)  -> prioridade 1
  * "tomad" ou "contrat" (valor tomado, valor contrato/contratado) -> prioridade 2
  * "limite" (limite aprovado/total)                               -> prioridade 3
- Se não achar NENHUMA dessas colunas -> "Limite Utilizado" = "A confirmar"

Uso (PowerShell):
  .\.venv\Scripts\python.exe .\02_SCRIPTS\agente_03_endividamento.py "C:\caminho\arquivo.pdf" --out "C:\saida\ENDIVIDAMENTO_EMPRESA.xlsx"

Dependências (pip):
  pip install pandas openpyxl pdfplumber python-docx pytesseract opencv-python pillow numpy
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd

import pdfplumber
from docx import Document

import cv2
import pytesseract


OUTPUT_COLUMNS = [
    "Instituição Financeira",
    "Tipo",
    "Linha de Crédito",
    "Limite Utilizado",
    "CNPJ",
]

BANK_KEYWORDS = [
    "banco do brasil", "bb", "b brasil", "bbrasil",
    "caixa", "bradesco", "itau", "santander", "safra", "btg", "daycoval",
    "sicoob", "sicredi", "inter", "original", "banrisul", "bv", "votorantim",
    "c6", "bmg", "nubank", "alfa", "abc", "modal", "pan", "omni", "brde",
]

INST_KEYS = ["institui", "banco", "credor", "instituicao financeira"]
LINE_KEYS = ["linha", "modal", "produto", "descr", "opera", "operacao", "carteira", "tipo"]
VALUE_KEYS_ANY = ["utiliz", "tomad", "contrat", "limite", "valor", "saldo", "devedor", "financiad"]
VALUE_PRIORITIES = ["utilizado", "tomado", "saldo", "financiado", "limite", "valor"]


def _norm(s: object) -> str:
    if s is None:
        return ""
    s2 = str(s).strip()
    s2 = unicodedata.normalize("NFKD", s2)
    s2 = "".join(ch for ch in s2 if not unicodedata.combining(ch))
    s2 = re.sub(r"\s+", " ", s2)
    return s2.lower().strip()


def _clean_spaces(s: object) -> str:
    if s is None:
        return ""
    return re.sub(r"\s+", " ", str(s)).strip()


def _is_na(v: object) -> bool:
    try:
        return bool(pd.isna(v))
    except Exception:
        return False


def find_cnpj_anywhere(text: str) -> Optional[str]:
    if not text:
        return None

    m = re.search(r"\b\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}\b", text)
    if m:
        return m.group(0)

    digits = re.sub(r"\D", "", text)
    m2 = re.search(r"\d{14}", digits)
    if m2:
        d = m2.group(0)
        return f"{d[0:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:14]}"
    return None


def classify_tipo(inst_name: str) -> str:
    s = _norm(inst_name)
    for k in BANK_KEYWORDS:
        if k in s:
            return "Banco"
    return "Fundo"


def parse_money_to_ptbr_text(val: object) -> Optional[str]:
    if val is None or _is_na(val):
        return None

    s = str(val).strip()
    if not s:
        return None

    s = s.replace("R$", "").replace("r$", "").strip()
    s = s.replace("\u00a0", " ").strip()

    digits = re.sub(r"[^\d\.,\-]", "", s)
    if not digits:
        return None

    neg = "-" in digits
    digits = digits.replace("-", "")
    if digits.count(",") > 1 and digits.rfind(",") >= 0:
        last = digits.rfind(",")
        digits = re.sub(r"[,\.\s]", "", digits[:last]) + "." + re.sub(r"[^\d]", "", digits[last + 1 :])
    elif digits.count(".") > 1 and digits.rfind(".") >= 0:
        last = digits.rfind(".")
        digits = re.sub(r"[,\.\s]", "", digits[:last]) + "." + re.sub(r"[^\d]", "", digits[last + 1 :])

    num: Optional[float] = None
    try:
        if "," in digits and "." in digits:
            num = float(digits.replace(".", "").replace(",", "."))
        elif "," in digits:
            parts = digits.split(",")
            if len(parts[-1]) == 2:
                num = float(digits.replace(".", "").replace(",", "."))
            else:
                num = float(digits.replace(",", ""))
        elif "." in digits:
            parts = digits.split(".")
            if len(parts[-1]) == 2:
                num = float(digits.replace(",", ""))
            else:
                num = float(digits.replace(".", ""))
        else:
            num = float(digits)

        if neg and num is not None:
            num = -num
    except Exception:
        return None

    if num is None:
        return None

    sign = "-" if num < 0 else ""
    num = abs(num)

    inteiro = int(num)
    cent = int(round((num - inteiro) * 100))
    if cent == 100:
        inteiro += 1
        cent = 0

    inteiro_str = f"{inteiro:,}".replace(",", ".")
    return f"{sign}{inteiro_str},{cent:02d}"


def detect_header_row(df_raw: pd.DataFrame, max_rows: int = 30, min_hits: int = 2) -> Optional[int]:
    best_i = None
    best_hits = 0

    n = min(max_rows, len(df_raw))
    for i in range(n):
        row = df_raw.iloc[i].tolist()
        cells = [_norm(c) for c in row]

        hits = 0
        if any(any(k in c for k in INST_KEYS) for c in cells):
            hits += 1
        if any(any(k in c for k in LINE_KEYS) for c in cells):
            hits += 1
        if any(any(k in c for k in VALUE_KEYS_ANY) for c in cells):
            hits += 1

        if hits > best_hits:
            best_hits = hits
            best_i = i

    if best_hits >= min_hits:
        return best_i
    return None


def table_is_valid(headers: List[str]) -> bool:
    cells = [_norm(h) for h in headers]
    hits = 0
    if any(any(k in c for k in INST_KEYS) for c in cells):
        hits += 1
    if any(any(k in c for k in LINE_KEYS) for c in cells):
        hits += 1
    if any(any(k in c for k in VALUE_KEYS_ANY) for c in cells):
        hits += 1
    return hits >= 2


def read_xlsx_or_csv_table(path: Path) -> Tuple[pd.DataFrame, str]:
    ext = path.suffix.lower()

    if ext in [".xlsx", ".xls"]:
        raw = pd.read_excel(path, header=None, dtype=str)
        header_row = detect_header_row(raw) or 0
        headers = [str(h).strip() if h is not None else "" for h in raw.iloc[header_row].tolist()]
        df = raw.iloc[header_row + 1 :].copy()
        df.columns = headers
        df = df.dropna(axis=1, how="all").dropna(axis=0, how="all")
        full_text = " ".join([_clean_spaces(x) for x in raw.astype(str).fillna("").values.flatten().tolist()])
        return df, full_text

    if ext == ".csv":
        df = pd.read_csv(path, dtype=str, encoding="utf-8", sep=None, engine="python")
        full_text = " ".join([_clean_spaces(x) for x in df.astype(str).fillna("").values.flatten().tolist()])
        return df, full_text

    raise ValueError(f"Formato não suportado: {ext}")


def read_pdf_first_valid_table(path: Path, max_tables: int = 5, max_pages: int = 10) -> Tuple[pd.DataFrame, str]:
    with pdfplumber.open(str(path)) as pdf:
        text = "\n".join([(p.extract_text() or "") for p in pdf.pages[: min(3, len(pdf.pages))]])

        found = 0
        for page in pdf.pages[:max_pages]:
            for tab in (page.extract_tables() or []):
                if not tab or len(tab) < 2:
                    continue
                found += 1
                if found > max_tables:
                    break

                best_i = None
                best_hits = 0
                for i in range(min(3, len(tab))):
                    headers = [(tab[i][j] or "").strip() for j in range(len(tab[i]))]
                    if not headers:
                        continue
                    if table_is_valid(headers):
                        cells = [_norm(h) for h in headers]
                        hits = 0
                        if any(any(k in c for k in INST_KEYS) for c in cells): hits += 1
                        if any(any(k in c for k in LINE_KEYS) for c in cells): hits += 1
                        if any(any(k in c for k in VALUE_KEYS_ANY) for c in cells): hits += 1
                        if hits > best_hits:
                            best_hits = hits
                            best_i = i

                if best_i is None:
                    continue

                headers = [(tab[best_i][j] or "").strip() for j in range(len(tab[best_i]))]
                data = tab[best_i + 1 :]
                df = pd.DataFrame(data, columns=headers).dropna(axis=0, how="all").dropna(axis=1, how="all")
                return df, text

            if found > max_tables:
                break

    raise RuntimeError("Não encontrei tabela válida no PDF (nas primeiras páginas/tabelas).")


def _find_pdftoppm(poppler_bin: Optional[Path]) -> Optional[Path]:
    if poppler_bin:
        exe = poppler_bin / "pdftoppm.exe"
        if exe.exists():
            return exe
    found = shutil.which("pdftoppm")
    return Path(found) if found else None


def read_pdf_table_ocr(path: Path, poppler_bin: Optional[Path]) -> Tuple[pd.DataFrame, str]:
    pdftoppm = _find_pdftoppm(poppler_bin)
    if not pdftoppm:
        return pd.DataFrame(), ""

    full_text_parts: List[str] = []
    with tempfile.TemporaryDirectory() as tmp_dir:
        prefix = Path(tmp_dir) / "page"
        cmd = [
            str(pdftoppm),
            "-f",
            "1",
            "-l",
            "5",
            "-png",
            str(path),
            str(prefix),
        ]
        try:
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        except Exception:
            return pd.DataFrame(), ""

        images = sorted(Path(tmp_dir).glob("page-*.png"))
        best_df = pd.DataFrame()
        for img_path in images:
            df, text = read_image_table_ocr(img_path)
            if text:
                full_text_parts.append(text)
            if not df.empty and table_is_valid(list(df.columns)):
                return df, "\n".join(full_text_parts)
            if best_df.empty and not df.empty:
                best_df = df

    return best_df, "\n".join(full_text_parts)


def read_docx_first_valid_table(path: Path, max_tables: int = 5) -> Tuple[pd.DataFrame, str]:
    doc = Document(str(path))
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text])

    checked = 0
    for t in doc.tables:
        checked += 1
        if checked > max_tables:
            break

        rows = []
        for row in t.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        if not rows or len(rows) < 2:
            continue

        headers = rows[0]
        if not table_is_valid(headers):
            continue

        df = pd.DataFrame(rows[1:], columns=headers).dropna(axis=0, how="all").dropna(axis=1, how="all")
        return df, full_text

    raise RuntimeError("Não encontrei tabela válida no DOCX (nas primeiras tabelas).")


def read_image_table_ocr(path: Path) -> Tuple[pd.DataFrame, str]:
    img = cv2.imread(str(path))
    if img is None:
        raise RuntimeError("Não consegui abrir a imagem para OCR.")

    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    gray = cv2.resize(gray, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
    gray = cv2.GaussianBlur(gray, (3, 3), 0)
    th = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 41, 15)

    text = pytesseract.image_to_string(
        th, config="--oem 3 --psm 6 -c preserve_interword_spaces=1"
    ) or ""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    header_idx = None
    best_hits = 0
    for i, ln in enumerate(lines[:30]):
        ln_norm = _norm(ln)
        hits = 0
        if any(k in ln_norm for k in INST_KEYS): hits += 1
        if any(k in ln_norm for k in LINE_KEYS): hits += 1
        if any(k in ln_norm for k in VALUE_KEYS_ANY): hits += 1
        if hits > best_hits:
            best_hits = hits
            header_idx = i

    if header_idx is None or best_hits < 2:
        return pd.DataFrame(), text

    header_parts = re.split(r"\s{2,}", lines[header_idx])
    header_parts = [h.strip() for h in header_parts if h.strip()]

    data_rows = []
    for ln in lines[header_idx + 1 :]:
        parts = re.split(r"\s{2,}", ln)
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) < 2:
            continue
        if len(parts) < len(header_parts):
            parts += [""] * (len(header_parts) - len(parts))
        if len(parts) > len(header_parts):
            parts = parts[: len(header_parts)]
        data_rows.append(parts)

    df = pd.DataFrame(data_rows, columns=header_parts).dropna(axis=0, how="all").dropna(axis=1, how="all")
    return df, text


def categorize_value_col(colname: str) -> Optional[str]:
    c = _norm(colname)
    if "utiliz" in c:
        return "utilizado"
    if "vencer" in c:
        return "utilizado"
    if "tomad" in c:
        return "tomado"
    if "contrat" in c:
        return "tomado"
    if "financi" in c:
        return "financiado"
    if "saldo" in c or "devedor" in c:
        return "saldo"
    if "valor" in c:
        return "valor"
    if "limite" in c:
        return "limite"
    return None


def pick_best_columns(df: pd.DataFrame) -> Tuple[Optional[str], Optional[str], Dict[str, List[str]]]:
    cols = list(df.columns)
    inst_col = next((c for c in cols if any(k in _norm(c) for k in INST_KEYS)), None)
    line_col = next((c for c in cols if any(k in _norm(c) for k in LINE_KEYS)), None)

    groups: Dict[str, List[str]] = {
        "utilizado": [],
        "tomado": [],
        "saldo": [],
        "financiado": [],
        "limite": [],
        "valor": [],
    }
    for c in cols:
        cat = categorize_value_col(str(c))
        if cat:
            groups[cat].append(c)

    return inst_col, line_col, groups


def choose_value_from_row(row: Dict[str, object], groups: Dict[str, List[str]]) -> str:
    for v in row.values():
        if v is None or _is_na(v):
            continue
        s = _clean_spaces(v)
        if not s:
            continue
        if "utiliz" in _norm(s):
            txt = parse_money_to_ptbr_text(s)
            if txt is not None:
                return txt

    for cat in VALUE_PRIORITIES:
        for c in groups.get(cat, []):
            txt = parse_money_to_ptbr_text(row.get(c))
            if txt is not None:
                return txt

    for v in row.values():
        if v is None or _is_na(v):
            continue
        txt = parse_money_to_ptbr_text(v)
        if txt is not None:
            return txt
    return "A confirmar"


def build_output(df: pd.DataFrame, full_text_for_cnpj: str, cnpj_override: str = "") -> pd.DataFrame:
    inst_col, line_col, groups = pick_best_columns(df)
    cnpj = cnpj_override.strip() if cnpj_override else (find_cnpj_anywhere(full_text_for_cnpj) or "A confirmar")

    rows = []
    for _, r in df.iterrows():
        inst = r.get(inst_col) if inst_col else None
        linha = r.get(line_col) if line_col else None

        inst = None if inst is None or _is_na(inst) or _norm(inst) in ["", "nan", "none"] else _clean_spaces(inst)
        linha = None if linha is None or _is_na(linha) or _norm(linha) in ["", "nan", "none"] else _clean_spaces(linha)

        if not inst and not linha:
            continue

        inst_out = inst or "A confirmar"
        linha_out = linha or "A confirmar"
        valor = choose_value_from_row(r.to_dict(), groups)

        rows.append({
            "Instituição Financeira": inst_out,
            "Tipo": classify_tipo(inst_out),
            "Linha de Crédito": linha_out,
            "Limite Utilizado": valor,
            "CNPJ": cnpj,
        })

    out = pd.DataFrame(rows, columns=OUTPUT_COLUMNS)

    if not out.empty:
        out["_k"] = (
            out["Instituição Financeira"].astype(str).str.lower().str.strip()
            + "|"
            + out["Linha de Crédito"].astype(str).str.lower().str.strip()
            + "|"
            + out["Limite Utilizado"].astype(str).str.strip()
        )
        out = out.drop_duplicates("_k").drop(columns=["_k"])
        out = out.sort_values("Instituição Financeira", kind="stable").reset_index(drop=True)

    return out


def read_any_table(path: Path, poppler_bin: Optional[Path] = None) -> Tuple[pd.DataFrame, str]:
    ext = path.suffix.lower()
    if ext in [".xlsx", ".xls", ".csv"]:
        return read_xlsx_or_csv_table(path)
    if ext == ".pdf":
        try:
            return read_pdf_first_valid_table(path)
        except RuntimeError:
            df, text = read_pdf_table_ocr(path, poppler_bin)
            if not df.empty:
                return df, text
            raise
    if ext == ".docx":
        return read_docx_first_valid_table(path)
    if ext in [".png", ".jpg", ".jpeg"]:
        return read_image_table_ocr(path)
    raise ValueError(f"Formato não suportado: {ext}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input", help="Arquivo de endividamento (xlsx/csv/pdf/docx/png/jpg)")
    ap.add_argument("--out", default="", help="Arquivo .xlsx de saida. Se vazio, gera em 03_OUTPUT\\3. Endividamento.")
    ap.add_argument("--cnpj", default="", help="Override manual do CNPJ (99.999.999/9999-99)")
    ap.add_argument("--tesseract", default="", help="Caminho do tesseract.exe (opcional)")
    ap.add_argument("--poppler", default="", help="Pasta do poppler/bin (opcional, para OCR em PDF)")
    ap.add_argument("--sheet", default="Planilha1", help="Nome da aba (default Planilha1)")
    args = ap.parse_args()

    in_path = Path(args.input)
    if not in_path.exists():
        raise SystemExit(f"Arquivo não encontrado: {in_path}")

    if args.tesseract:
        pytesseract.pytesseract.tesseract_cmd = args.tesseract

    poppler_bin = None
    poppler_arg = args.poppler.strip() if args.poppler else os.environ.get("POPPLER_BIN", "")
    if poppler_arg:
        poppler_bin = Path(poppler_arg)

    df, full_text = read_any_table(in_path, poppler_bin=poppler_bin)

    if df.empty:
        out_df = pd.DataFrame(columns=OUTPUT_COLUMNS)
    else:
        out_df = build_output(df, full_text_for_cnpj=full_text, cnpj_override=args.cnpj)

    if args.out:
        out_path = Path(args.out)
    else:
        base_dir = Path(__file__).resolve().parents[1]
        out_dir = base_dir / "03_OUTPUT" / "3. Endividamento"
        out_path = out_dir / f"ENDIVIDAMENTO_{in_path.stem}.xlsx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
        out_df.to_excel(writer, sheet_name=args.sheet, index=False)

    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()
